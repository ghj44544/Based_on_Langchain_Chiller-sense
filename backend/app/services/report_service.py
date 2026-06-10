import json
import re
from datetime import datetime
from html import escape
from pathlib import Path

from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.db.models import DiagnosisRecord
from app.services.preprocess_service import load_model_feature_columns
from app.utils.file_utils import ensure_dir


class ReportNotFoundError(ValueError):
    pass


class ReportFileError(ValueError):
    pass


class ReportFormatError(ValueError):
    pass


SUPPORTED_REPORT_FORMATS = {"md", "html", "docx"}
REPORT_MEDIA_TYPES = {
    "md": "text/markdown; charset=utf-8",
    "html": "text/html; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


SEVERITY_TEXT = {
    "high": "高：主导类别集中度较高，建议优先安排现场核查。",
    "medium": "中：存在一定类别集中趋势，建议结合工况和历史数据复核。",
    "low": "低：预测分布较分散，本报告更适合作为整体分布参考。",
}


FAULT_SUGGESTIONS = {
    "制冷剂过量": [
        "核对制冷剂充注量，检查过冷度、冷凝压力和压缩机负荷是否偏高。",
        "必要时按规范回收多余制冷剂，并复测 COP 与 kW/Ton。",
    ],
    "制冷剂泄漏": [
        "检查管路、阀件、换热器和连接处是否存在泄漏点。",
        "修复泄漏后抽真空并按标准重新充注制冷剂。",
    ],
    "润滑油加注过量": [
        "检查油位、油压和油分离器状态，确认是否存在油循环异常。",
        "按厂家规范调整油量，维修后观察压缩机运行电流和排气状态。",
    ],
    "不凝性气体": [
        "检查系统密封性和抽真空质量，关注冷凝压力是否异常升高。",
        "按维护规程排除不凝性气体，并复核冷凝器换热状态。",
    ],
    "冷凝器结垢": [
        "检查冷凝器管束、水质和结垢情况，必要时安排清洗。",
        "复核冷却水进出水温差、冷凝压力和 kW/Ton 的变化。",
    ],
    "冷却水流量不足": [
        "检查冷却水泵、阀门开度、过滤器、管路阻塞和冷却塔运行状态。",
        "复核 FWC、冷凝器进出水温差、冷凝压力和机组能效。",
    ],
    "冷冻水流量不足": [
        "检查冷冻水泵、阀门、过滤器和管路阻塞情况。",
        "复核 FWE、蒸发器水侧温差、蒸发压力和出水温度稳定性。",
    ],
    "正常": [
        "当前主导类别为正常时，仍建议持续监测 COP、kW/Ton 和水侧温差。",
        "若出现局部类别异常升高，可截取对应时段单独复诊。",
    ],
}


def _load_label_map() -> dict[str, str]:
    path = get_settings().label_map_file
    return json.loads(path.read_text(encoding="utf-8"))


def _percent(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{value * 100:.2f}%"


def _format_confidence(confidence: float | None) -> str:
    if confidence is None:
        return "模型未提供概率置信度。"
    if confidence >= 0.85:
        return "平均置信度较高，模型对样本级预测较稳定。"
    if confidence >= 0.6:
        return "平均置信度中等，建议结合现场工况复核。"
    return "平均置信度偏低，建议优先将本报告作为筛查参考，而不是最终维修结论。"


def _dataset_interpretation(record: DiagnosisRecord, label_distribution: dict[str, int]) -> str:
    if not label_distribution:
        return "上传数据未包含 label，报告按待诊断数据处理。"

    unique_count = len(label_distribution)
    if unique_count > 1:
        return (
            f"上传数据包含 {unique_count} 个真实标签类别，说明它更接近测试/评估数据集。"
            "此时主导类别代表整批样本的多数投票结果，不宜直接解释为单台机组当前唯一故障。"
        )
    return "上传数据只包含单一真实标签类别，更适合生成某一故障工况的诊断报告。"


def _distribution_table(distribution: dict[str, int], total_rows: int) -> str:
    label_map = _load_label_map()
    lines = [
        "| 标签 | 故障名称 | 样本数 | 占比 |",
        "| --- | --- | ---: | ---: |",
    ]
    for label, count in sorted(distribution.items(), key=lambda item: int(item[0])):
        ratio = count / total_rows if total_rows else 0
        lines.append(
            f"| {label} | {label_map.get(str(label), '未知类别')} | {count} | {_percent(ratio)} |"
        )
    return "\n".join(lines)


def _label_distribution_table(label_distribution: dict[str, int], total_rows: int) -> str:
    if not label_distribution:
        return "上传数据未包含 label。"

    label_map = _load_label_map()
    lines = [
        "| 真实标签 | 故障名称 | 样本数 | 占比 |",
        "| --- | --- | ---: | ---: |",
    ]
    for label, count in sorted(label_distribution.items(), key=lambda item: int(item[0])):
        ratio = count / total_rows if total_rows else 0
        lines.append(
            f"| {label} | {label_map.get(str(label), '未知类别')} | {count} | {_percent(ratio)} |"
        )
    return "\n".join(lines)


def _dedupe(items: list[str]) -> list[str]:
    result = []
    seen = set()
    for item in items:
        normalized = item.strip()
        if "ENABLE_LLM" in normalized or "label_map.json" in normalized:
            continue
        if normalized and normalized not in seen:
            result.append(normalized)
            seen.add(normalized)
    return result


def _fault_specific_suggestions(fault_name: str) -> list[str]:
    for key, suggestions in FAULT_SUGGESTIONS.items():
        if key in fault_name:
            return suggestions
    return []


def _monitoring_suggestions(fault_name: str) -> list[str]:
    base = [
        "持续观察 COP、kW/Ton、冷冻水/冷却水进出水温度和流量变化。",
        "维修或调整后重新上传同一工况数据，比较主导类别、平均置信度和预测分布变化。",
    ]
    if "冷却水流量不足" in fault_name:
        base.insert(0, "重点跟踪 FWC、TWCI、TWCO、冷凝压力 PRC 和冷凝温度相关指标。")
    elif "冷冻水流量不足" in fault_name:
        base.insert(0, "重点跟踪 FWE、TWEI、TWEO、蒸发器负荷和出水温度稳定性。")
    elif "制冷剂" in fault_name:
        base.insert(0, "重点跟踪 PRC、TRC、TRC_sub、P_lift、kW/Ton 和 COP。")
    return base


def _format_list(items: list[str]) -> str:
    if not items:
        return "暂无。"
    return "\n".join(f"- {item}" for item in items)


def _model_feature_text() -> str:
    try:
        features = load_model_feature_columns()
    except Exception:
        features = []
    if not features:
        return "未配置单独的模型输入特征，默认使用完整 RP1043 特征。"
    return "、".join(features)


def _validate_report_format(report_format: str) -> str:
    normalized = report_format.lower().strip().lstrip(".")
    if normalized not in SUPPORTED_REPORT_FORMATS:
        raise ReportFormatError("报告格式仅支持 md、html、docx")
    return normalized


def _build_report_markdown(record: DiagnosisRecord) -> str:
    model_result = json.loads(record.model_result_json or "{}")
    distribution = json.loads(record.prediction_distribution_json or "{}")
    label_distribution = json.loads(record.label_distribution_json or "{}")
    suggestions = json.loads(record.maintenance_suggestions_json or "[]")
    distribution = {str(key): int(value) for key, value in distribution.items()}
    label_distribution = {str(key): int(value) for key, value in label_distribution.items()}

    generated_at = datetime.utcnow().isoformat() + "Z"
    dominant_ratio = record.dominant_ratio or 0
    avg_confidence = record.avg_confidence
    severity = model_result.get("severity", "unknown")
    dataset_note = _dataset_interpretation(record, label_distribution)
    confidence_note = _format_confidence(avg_confidence)
    severity_note = SEVERITY_TEXT.get(severity, "未知：请结合预测分布和现场工况判断。")
    fault_suggestions = _fault_specific_suggestions(record.dominant_fault_name or "")
    maintenance_suggestions = _dedupe(fault_suggestions + suggestions)
    monitoring = _monitoring_suggestions(record.dominant_fault_name or "")

    return f"""# 冷水机组故障诊断报告

## 一、诊断摘要

| 项目 | 结果 |
| --- | --- |
| 诊断记录 ID | {record.id} |
| 主导标签 | {record.dominant_label} |
| 主导故障名称 | {record.dominant_fault_name} |
| 主导比例 | {_percent(dominant_ratio)} |
| 平均置信度 | {_percent(avg_confidence) if avg_confidence is not None else "未提供"} |
| 严重程度 | {severity} |

结论：本批数据的多数投票类别为 **{record.dominant_fault_name}**。{severity_note}

可信度解读：{confidence_note}

数据解读：{dataset_note}

## 二、数据基本信息

- 文件名：{record.filename}
- 样本数：{record.total_rows}
- 字段数：{record.total_columns}
- 是否包含 label：{record.has_label}
- 模型输入特征：{_model_feature_text()}

### 真实标签分布

{_label_distribution_table(label_distribution, record.total_rows)}

## 三、模型诊断结果

- 主导标签：{record.dominant_label}
- 主导故障名称：{record.dominant_fault_name}
- 主导比例：{record.dominant_ratio}（{_percent(dominant_ratio)}）
- 平均置信度：{record.avg_confidence if record.avg_confidence is not None else "未提供"}
- 严重程度：{severity}

## 四、预测分布

{_distribution_table(distribution, record.total_rows)}

## 五、故障解释

{record.explanation or "暂无解释。"}

## 六、维修建议

{_format_list(maintenance_suggestions)}

## 七、后续监测建议

{_format_list(monitoring)}

## 八、报告备注

- 模型预测结果来自 MATLAB ultra 模型；LangChain 仅用于解释和问答增强，不改变模型输出。
- 如果上传数据同时包含多个真实故障类别，本报告应作为批量评估/分布分析，而不是单一故障工单。
- 若需要更完整的自然语言解释，可在 `.env` 中启用 `ENABLE_LLM=true` 并配置模型 API。

## 九、生成时间

{generated_at}
"""


def _inline_markdown(text: str) -> str:
    escaped = escape(text)
    return re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)


def _markdown_to_html(markdown_text: str) -> str:
    lines = markdown_text.splitlines()
    html_parts = [
        "<!doctype html>",
        '<html lang="zh-CN">',
        "<head>",
        '<meta charset="utf-8">',
        "<title>冷水机组故障诊断报告</title>",
        "<style>",
        "body{font-family:Arial,'Microsoft YaHei',sans-serif;line-height:1.7;color:#1f2937;max-width:960px;margin:32px auto;padding:0 24px;}",
        "h1{font-size:28px;border-bottom:2px solid #2563eb;padding-bottom:12px;}",
        "h2{font-size:22px;margin-top:32px;color:#1d4ed8;}",
        "h3{font-size:18px;margin-top:24px;color:#334155;}",
        "table{border-collapse:collapse;width:100%;margin:12px 0 20px;}",
        "th,td{border:1px solid #cbd5e1;padding:8px 10px;text-align:left;}",
        "th{background:#eff6ff;}",
        "li{margin:4px 0;}",
        "p{margin:10px 0;}",
        "</style>",
        "</head><body>",
    ]

    in_ul = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            if in_ul:
                html_parts.append("</ul>")
                in_ul = False
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("| ---"):
            if in_ul:
                html_parts.append("</ul>")
                in_ul = False
            headers = [cell.strip() for cell in line.strip("|").split("|")]
            html_parts.append("<table><thead><tr>")
            html_parts.extend(f"<th>{escape(header)}</th>" for header in headers)
            html_parts.append("</tr></thead><tbody>")
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                html_parts.append("<tr>")
                html_parts.extend(f"<td>{_inline_markdown(cell)}</td>" for cell in cells)
                html_parts.append("</tr>")
                i += 1
            html_parts.append("</tbody></table>")
            continue

        if line.startswith("# "):
            if in_ul:
                html_parts.append("</ul>")
                in_ul = False
            html_parts.append(f"<h1>{_inline_markdown(line[2:].strip())}</h1>")
        elif line.startswith("## "):
            if in_ul:
                html_parts.append("</ul>")
                in_ul = False
            html_parts.append(f"<h2>{_inline_markdown(line[3:].strip())}</h2>")
        elif line.startswith("### "):
            if in_ul:
                html_parts.append("</ul>")
                in_ul = False
            html_parts.append(f"<h3>{_inline_markdown(line[4:].strip())}</h3>")
        elif line.startswith("- "):
            if not in_ul:
                html_parts.append("<ul>")
                in_ul = True
            html_parts.append(f"<li>{_inline_markdown(line[2:].strip())}</li>")
        else:
            if in_ul:
                html_parts.append("</ul>")
                in_ul = False
            html_parts.append(f"<p>{_inline_markdown(line)}</p>")
        i += 1

    if in_ul:
        html_parts.append("</ul>")
    html_parts.append("</body></html>")
    return "\n".join(html_parts)


def _markdown_to_docx(markdown_text: str, output_path: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise ReportFormatError("导出 DOCX 需要安装 python-docx，请执行 pip install python-docx。") from exc

    document = Document()
    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("| ---"):
            headers = [cell.strip() for cell in line.strip("|").split("|")]
            rows = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([cell.strip() for cell in lines[i].strip().strip("|").split("|")])
                i += 1
            table = document.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for idx, header in enumerate(headers):
                table.rows[0].cells[idx].text = header
            for row in rows:
                cells = table.add_row().cells
                for idx, cell in enumerate(row[: len(headers)]):
                    cells[idx].text = cell.replace("**", "")
            continue

        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip().replace("**", ""), style="List Bullet")
        else:
            document.add_paragraph(line.replace("**", ""))
        i += 1

    document.save(output_path)


def generate_report_file(diagnosis_id: int, db: Session, report_format: str = "md") -> Path:
    report_format = _validate_report_format(report_format)
    record = db.get(DiagnosisRecord, diagnosis_id)
    if record is None:
        raise ReportNotFoundError("诊断记录不存在")

    settings = get_settings()
    report_dir = ensure_dir(settings.report_path)
    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    report_path = report_dir / f"diagnosis_report_{diagnosis_id}_{timestamp}.{report_format}"
    markdown_text = _build_report_markdown(record)

    if report_format == "md":
        report_path.write_text(markdown_text, encoding="utf-8")
    elif report_format == "html":
        report_path.write_text(_markdown_to_html(markdown_text), encoding="utf-8")
    elif report_format == "docx":
        _markdown_to_docx(markdown_text, report_path)
    return report_path


def generate_markdown_report(diagnosis_id: int, db: Session) -> Path:
    return generate_report_file(diagnosis_id, db, report_format="md")


def get_report_file(filename: str) -> Path:
    safe_name = Path(filename).name
    suffix = Path(safe_name).suffix.lower().lstrip(".")
    if safe_name != filename or suffix not in SUPPORTED_REPORT_FORMATS:
        raise ReportFileError("报告文件名不合法")

    report_dir = ensure_dir(get_settings().report_path).resolve()
    report_path = (report_dir / safe_name).resolve()
    if report_path.parent != report_dir:
        raise ReportFileError("报告文件路径不合法")
    if not report_path.exists():
        raise ReportNotFoundError("报告文件不存在")
    return report_path


def find_latest_report(diagnosis_id: int, report_format: str | None = None) -> Path:
    report_dir = ensure_dir(get_settings().report_path)
    suffix = "*"
    if report_format:
        suffix = _validate_report_format(report_format)
    reports = sorted(
        report_dir.glob(f"diagnosis_report_{diagnosis_id}_*.{suffix}"),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    if not reports:
        raise ReportNotFoundError("该诊断记录还没有生成报告")
    return reports[0]


def build_report_export_data(report_path: Path) -> dict[str, str]:
    return {
        "report_path": str(report_path),
        "report_filename": report_path.name,
        "report_format": report_path.suffix.lower().lstrip("."),
        "download_url": f"/api/report/download/{report_path.name}",
    }


def get_report_media_type(report_path: Path) -> str:
    return REPORT_MEDIA_TYPES.get(report_path.suffix.lower().lstrip("."), "application/octet-stream")
